#!/usr/bin/env python3
"""
intersperse_payloads.py

Randomly intersperse payload values from a CSV into DOCX and CSV files.

For CSV targets:  payload rows are inserted at random positions. Columns are matched
                  by name where possible; unmatched payload values are concatenated
                  into the first column of the target.

For DOCX targets: each payload row is converted to a text string (values joined with " | ")
                  and inserted as a paragraph at a random position in the document.

Usage:
    python intersperse_payloads.py \\
        --input-folder /path/to/files \\
        --payload-csv /path/to/payloads.csv \\
        --output-folder /path/to/output \\
        [--count N] \\
        [--seed 42]

    # Generate a sample payload CSV to get started:
    python intersperse_payloads.py --generate-sample-payload sample_payload.csv
"""

import argparse
import csv
import os
import random
import shutil
from pathlib import Path

try:
    from docx import Document
    from docx.oxml import OxmlElement
except ImportError:
    raise SystemExit(
        "Missing dependency: pip install python-docx"
    )


# ---------------------------------------------------------------------------
# Argument parsing
# ---------------------------------------------------------------------------

def parse_args():
    parser = argparse.ArgumentParser(
        description="Randomly intersperse payload values into DOCX and CSV files.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument(
        "--input-folder",
        help="Folder containing .docx and .csv files to modify.",
    )
    parser.add_argument(
        "--payload-csv",
        help="CSV file whose rows will be injected into the target files.",
    )
    parser.add_argument(
        "--output-folder",
        default="./output",
        help="Destination folder for modified files (default: ./output).",
    )
    parser.add_argument(
        "--count",
        type=int,
        default=None,
        help="Number of payload items to inject per file. "
             "Defaults to a random value between 3 and 10.",
    )
    parser.add_argument(
        "--seed",
        type=int,
        default=None,
        help="Random seed for reproducible output.",
    )
    parser.add_argument(
        "--generate-sample-payload",
        metavar="OUTPUT_CSV",
        help="Write a sample payload CSV to the given path and exit.",
    )
    return parser.parse_args()


# ---------------------------------------------------------------------------
# Payload utilities
# ---------------------------------------------------------------------------

def load_payload_csv(path: str) -> list[dict]:
    """Load every row of a CSV file into a list of dicts."""
    rows = []
    with open(path, newline="", encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        for row in reader:
            rows.append(dict(row))
    return rows


def payload_row_to_text(row: dict) -> str:
    """Flatten a payload row into a single readable string."""
    return " | ".join(str(v) for v in row.values() if str(v).strip())


def build_csv_row_from_payload(payload_row: dict, target_fieldnames: list[str]) -> dict:
    """
    Map a payload row onto the target CSV's column schema.

    Strategy:
      1. Copy values for columns whose names appear in both schemas.
      2. If nothing matched, concatenate all payload values and place them
         in the first target column so the injection is still visible.
    """
    result = {col: "" for col in target_fieldnames}
    matched_any = False

    for col in target_fieldnames:
        if col in payload_row:
            result[col] = payload_row[col]
            matched_any = True

    if not matched_any and target_fieldnames:
        result[target_fieldnames[0]] = payload_row_to_text(payload_row)

    return result


def pick_payloads(payload_rows: list[dict], count: int | None, rng: random.Random) -> list[dict]:
    """Choose *count* payload rows (with replacement) to inject."""
    n = count if count is not None else rng.randint(3, min(10, len(payload_rows)))
    n = min(n, len(payload_rows))
    return rng.choices(payload_rows, k=n)


# ---------------------------------------------------------------------------
# CSV processing
# ---------------------------------------------------------------------------

def intersperse_csv(
    input_path: str,
    output_path: str,
    payload_rows: list[dict],
    count: int | None,
    rng: random.Random,
) -> None:
    """Insert payload rows at random positions within a CSV file."""
    with open(input_path, newline="", encoding="utf-8") as fh:
        reader = csv.DictReader(fh)
        fieldnames: list[str] = list(reader.fieldnames or [])
        data_rows = list(reader)

    selected = pick_payloads(payload_rows, count, rng)
    n = len(selected)

    injected = [
        build_csv_row_from_payload(p, fieldnames) for p in selected
    ]

    # Insert each payload row at a random position (sequentially so indices remain valid).
    result = list(data_rows)
    for row in injected:
        pos = rng.randint(0, len(result))
        result.insert(pos, row)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    with open(output_path, "w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(result)

    print(f"  [CSV]  Injected {n} rows  → {output_path}")


# ---------------------------------------------------------------------------
# DOCX processing
# ---------------------------------------------------------------------------

def _make_paragraph_element(text: str):
    """Build a minimal w:p XML element containing the given text."""
    p = OxmlElement("w:p")
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    # Preserve leading/trailing whitespace
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p


def intersperse_docx(
    input_path: str,
    output_path: str,
    payload_rows: list[dict],
    count: int | None,
    rng: random.Random,
) -> None:
    """Insert payload paragraphs at random positions within a DOCX file."""
    doc = Document(input_path)

    num_paras = len(doc.paragraphs)
    if num_paras == 0:
        print(f"  [DOCX] No paragraphs found, copying unchanged → {output_path}")
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        shutil.copy2(input_path, output_path)
        return

    selected = pick_payloads(payload_rows, count, rng)
    n = len(selected)
    texts = [payload_row_to_text(row) for row in selected]

    # Determine insertion positions (each is an index into doc.paragraphs).
    # Process from highest to lowest so earlier indices remain stable
    # as we insert before later paragraphs.
    positions = sorted(
        [rng.randint(0, num_paras - 1) for _ in range(n)],
        reverse=True,
    )

    for text, pos in zip(texts, positions):
        # doc.paragraphs is re-evaluated from the live XML on every access.
        target_para_elem = doc.paragraphs[pos]._element
        new_p = _make_paragraph_element(text)
        target_para_elem.addprevious(new_p)

    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)
    print(f"  [DOCX] Injected {n} paragraphs → {output_path}")


# ---------------------------------------------------------------------------
# Sample payload generator
# ---------------------------------------------------------------------------

SAMPLE_PAYLOAD_ROWS = [
    {
        "payload_id": "P-001",
        "category": "auth",
        "description": "Unauthorised login attempt detected from external IP",
        "severity": "HIGH",
        "notes": "Reviewed by SOC",
    },
    {
        "payload_id": "P-002",
        "category": "data",
        "description": "Sensitive file accessed outside business hours",
        "severity": "MEDIUM",
        "notes": "Flagged for review",
    },
    {
        "payload_id": "P-003",
        "category": "network",
        "description": "Unexpected outbound connection to unknown endpoint",
        "severity": "HIGH",
        "notes": "Under investigation",
    },
    {
        "payload_id": "P-004",
        "category": "compliance",
        "description": "MFA bypass policy exception granted temporarily",
        "severity": "LOW",
        "notes": "Approved by CISO",
    },
    {
        "payload_id": "P-005",
        "category": "config",
        "description": "Security group rule opened port 22 to 0.0.0.0/0",
        "severity": "CRITICAL",
        "notes": "Remediated within SLA",
    },
]


def generate_sample_payload(output_path: str) -> None:
    fieldnames = list(SAMPLE_PAYLOAD_ROWS[0].keys())
    with open(output_path, "w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(SAMPLE_PAYLOAD_ROWS)
    print(f"Sample payload CSV written to: {output_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    args = parse_args()

    # --generate-sample-payload shortcut
    if args.generate_sample_payload:
        generate_sample_payload(args.generate_sample_payload)
        return

    if not args.input_folder or not args.payload_csv:
        raise SystemExit(
            "Error: --input-folder and --payload-csv are required.\n"
            "Run with --help for usage details."
        )

    rng = random.Random(args.seed)

    input_folder = Path(args.input_folder)
    output_folder = Path(args.output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    print(f"Loading payloads from: {args.payload_csv}")
    payload_rows = load_payload_csv(args.payload_csv)
    if not payload_rows:
        raise SystemExit("Error: payload CSV is empty.")
    print(f"  {len(payload_rows)} payload row(s) loaded.\n")

    all_files = sorted(input_folder.iterdir())
    csv_files = [f for f in all_files if f.suffix.lower() == ".csv"]
    docx_files = [f for f in all_files if f.suffix.lower() == ".docx"]

    print(
        f"Found {len(docx_files)} DOCX and {len(csv_files)} CSV file(s) "
        f"in: {input_folder}\n"
    )

    errors = []

    for f in csv_files:
        out = output_folder / f.name
        try:
            intersperse_csv(str(f), str(out), payload_rows, args.count, rng)
        except Exception as exc:
            errors.append((f.name, exc))
            print(f"  [CSV]  ERROR processing {f.name}: {exc}")

    for f in docx_files:
        out = output_folder / f.name
        try:
            intersperse_docx(str(f), str(out), payload_rows, args.count, rng)
        except Exception as exc:
            errors.append((f.name, exc))
            print(f"  [DOCX] ERROR processing {f.name}: {exc}")

    print(f"\nDone. Output written to: {output_folder}")
    if errors:
        print(f"  {len(errors)} file(s) had errors (see above).")


if __name__ == "__main__":
    main()
